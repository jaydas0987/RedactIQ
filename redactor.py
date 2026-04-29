#!/usr/bin/env python3
"""
PII Redactor — Offline, batch, three-mode
Supports: .txt, .pdf, .docx, .xlsx, .csv, .sql

Usage:
    python redactor.py /path/to/folder --mode public
    python redactor.py /path/to/folder --mode research
    python redactor.py /path/to/folder --mode audit

Modes:
    public    — fixed uniform blocks, no length information leaked
    research  — consistent pseudonyms per unique value (re-linkable, not re-identifiable)
    audit     — type labels only, financial/operational data kept intact

Dependencies:
    pip install pypdf python-docx openpyxl reportlab pdfplumber pikepdf spacy
    python -m spacy download en_core_web_lg
"""

import re
import sys
import csv
import json
import hashlib
import argparse
from pathlib import Path
from datetime import datetime

# NER detection layer (spaCy + Indian gazetteer + Stanza Hindi)
try:
    from ner_detect import ner_redact, load_spacy_model, print_setup
    NER_AVAILABLE = True
except ImportError:
    NER_AVAILABLE = False


# ─────────────────────────────────────────────────────────────
# TERMINAL COLOURS
# ─────────────────────────────────────────────────────────────

C = {
    "header": "\033[1;36m", "keep":   "\033[1;32m",
    "redact": "\033[1;31m", "warn":   "\033[1;33m",
    "dim":    "\033[2m",    "bold":   "\033[1m",
    "info":   "\033[1;34m", "reset":  "\033[0m",
}
def col(key, text): return f"{C.get(key,'')}{text}{C['reset']}"
def banner(text):   print(f"\n{col('header', '═'*60)}\n  {col('header', text)}\n{col('header', '═'*60)}")


# ─────────────────────────────────────────────────────────────
# REPLACEMENT STRATEGY
#
# PUBLIC   → fixed-width █ blocks per type — no length info leaked
# RESEARCH → consistent pseudonym token per unique value
#            same original value → same token across ALL files in batch
#            e.g. "Aarav Mehta" always becomes [NAME-R001]
# AUDIT    → type label only — [NAME], [EMAIL], [PHONE] etc.
# ─────────────────────────────────────────────────────────────

BLOCK = "█"

PUBLIC_FIXED = {
    "name":    "████████████",       # 12
    "email":   "████████████████",   # 16
    "phone":   "████████████",       # 12
    "ip":      "███████████",        # 11
    "aadhaar": "████████████",       # 12
    "pan":     "██████████",         # 10
    "ssn":     "███████████",        # 11
    "card":    "████████████████",   # 16
    "dob":     "██████████",         # 10
    "address": "████████████████",   # 16
    "generic": "████████████",       # 12
}

AUDIT_LABELS = {
    "name":    "[NAME]",
    "email":   "[EMAIL]",
    "phone":   "[PHONE]",
    "ip":      "[IP]",
    "aadhaar": "[AADHAAR]",
    "pan":     "[PAN]",
    "ssn":     "[SSN]",
    "card":    "[CARD]",
    "dob":     "[DOB]",
    "address": "[ADDRESS]",
    "upi":     "[UPI]",
    "generic": "[REDACTED]",
}


# ─────────────────────────────────────────────────────────────
# PSEUDONYM VAULT — research mode
# ─────────────────────────────────────────────────────────────

class PseudonymVault:
    """
    Maps each unique PII value to a consistent token within a batch run.
    Token is deterministic per (type, value) — same name always gets same token.
    Uses SHA-256 of the value as the internal key — original is never stored.
    Tokens reset between runs (in-memory only, nothing written to disk).
    """
    def __init__(self):
        self._counters = {}
        self._map = {}

    def get(self, pii_type: str, value: str) -> str:
        key = (pii_type, hashlib.sha256(value.lower().strip().encode()).hexdigest())
        if key not in self._map:
            n = self._counters.get(pii_type, 0) + 1
            self._counters[pii_type] = n
            self._map[key] = f"[{pii_type.upper()}-R{n:03d}]"
        return self._map[key]

    def summary(self):
        return dict(self._counters)


VAULT = PseudonymVault()


def replace(pii_type: str, value: str, mode: str) -> str:
    """Return the correct replacement for a PII value given the mode."""
    if mode == "public":
        return PUBLIC_FIXED.get(pii_type, PUBLIC_FIXED["generic"])
    elif mode == "research":
        return VAULT.get(pii_type, value)
    elif mode == "audit":
        return AUDIT_LABELS.get(pii_type, AUDIT_LABELS["generic"])
    return value


# ─────────────────────────────────────────────────────────────
# PII PATTERNS
# (pii_type, regex, redact_public, redact_research, redact_audit)
# ─────────────────────────────────────────────────────────────

PII_RULES = [
    # Card must come BEFORE aadhaar — 16-digit cards share the \d{4}\s\d{4}\s\d{4} prefix
    # and would be partially consumed by the aadhaar regex if it ran first.
    ("card",    r"\b(?:\d{4}[- ]){3}\d{4}\b",                                  True,  True,  True),
    ("aadhaar", r"\b\d{4}\s\d{4}\s\d{4}\b",                                   True,  True,  True),
    # PAN: allow optional unicode whitespace between alpha and digit blocks (malformed data)
    ("pan",     r"\b[A-Z]{5}[\s\u2800-\u28FF\u200B-\u200F]?\d{4}[A-Z]\b",     True,  True,  True),
    ("ssn",     r"\b\d{3}-\d{2}-\d{4}\b",                                      True,  True,  True),
    ("ip",      r"\b(?:\d{1,3}\.){3}\d{1,3}\b",                                True,  True,  True),
    ("email",   r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b",      True,  True,  True),
    # UPI: must come AFTER email — handles virtualaddress@bankhandle format without TLD
    ("upi",     r"\b[a-zA-Z0-9.\-_]{2,}@[a-zA-Z]{2,}\b",                      True,  True,  True),
    ("phone",   r"(?:\+91[-.\s]?)?\d{5}[-.\s]\d{5}"        # +91 98201 44312 (5+5)
                r"|(?:\+91[-.\s]?)?\d{5}\s?\d{6}"           # +91 98201 444312 (5+6)
                r"|\+91[-.\s]?\d{10}"                        # +91 9820144312 (no space)
                r"|\b(?:\+91[-.\s]?)?\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}\b"  # intl format
                r"|\b[6-9]\d{9}\b",                          # bare Indian mobile (no prefix)
                                                                True,  True,  True),
    ("dob",     r"\b(?:DOB|Date\s+of\s+Birth|Born)[:\s]+\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
                                                                                True,  False, True),
    # Address: number prefix + locality keyword (conservative, avoids false positives)
    ("address", r"\b(?:Flat|F\.?No\.?|Apartment|Apt\.?|Plot|Door\s+No\.?)"
                r"\s*[\w\-/]+(?:\s*,\s*[A-Za-z0-9\s\-]+)?"
                r"|\b\d{1,5}[A-Z]?\s+[A-Za-z0-9\s]{3,40}"
                r"(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|"
                r"Lane|Ln|Court|Ct|Way|Place|Pl|"
                r"Nagar|Marg|Vihar|Colony|Layout|Sector|Block|Extension|Ext|"
                r"Cross|Main|Circle|Chowk|Bazaar|Bazar|Gali|Mohalla)\.?\b",  True,  False, True),
    ("name",    r"\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+",
                                                                                True,  True,  True),
]


# ─────────────────────────────────────────────────────────────
# PII SCHEMA — user-classified unknown fields
# ─────────────────────────────────────────────────────────────

SCHEMA_FILE = "pii_schema.json"

def load_schema():
    if Path(SCHEMA_FILE).exists():
        with open(SCHEMA_FILE) as f:
            return json.load(f)
    return {}

def save_schema(schema):
    with open(SCHEMA_FILE, "w") as f:
        json.dump(schema, f, indent=2)


# ─────────────────────────────────────────────────────────────
# FILE READERS
# ─────────────────────────────────────────────────────────────

def read_txt(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_pdf(path):
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        print(col("warn", f"  [!] Could not read PDF {path.name}: {e}"))
        return ""

def read_docx(path):
    try:
        from docx import Document
        return "\n".join(p.text for p in Document(str(path)).paragraphs)
    except Exception as e:
        print(col("warn", f"  [!] Could not read DOCX {path.name}: {e}"))
        return ""

def read_csv(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        text = f.read()
    # Strip invisible unicode characters that appear in malformed PAN/Aadhaar data
    return re.sub(r'[⠀-⣿​-‏﻿ ]', ' ', text)

def read_xlsx(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                lines.append("\t".join(str(c) if c is not None else "" for c in row))
        return "\n".join(lines)
    except Exception as e:
        print(col("warn", f"  [!] Could not read XLSX {path.name}: {e}"))
        return ""

def read_sql(path):
    text = read_txt(path)
    # Collapse split string literals separated by SQL block comments
    # e.g. 'PQRSX' /*hidden*/ '9876Z' → 'PQRSX9876Z'
    text = re.sub(r"'\s*/\*.*?\*/\s*'", "", text, flags=re.DOTALL)
    return text
READERS = {".txt": read_txt, ".pdf": read_pdf, ".docx": read_docx,
           ".csv": read_csv, ".xlsx": read_xlsx, ".sql": read_sql}


# ─────────────────────────────────────────────────────────────
# AUTO-SCAN
# ─────────────────────────────────────────────────────────────

CUSTOMER_NAME_RE = re.compile(
    r"(?:"
    # Customer-facing keywords
    r"customer|user|client|referenced\s+customer|for\s+customer|linked\s+to\s+customer"
    r"|"
    # Staff/role keywords
    r"agent|officer|analyst|investigator|reviewer|auditor|examiner"
    r"|senior\s+analyst|reviewing\s+officer|fraud\s+team\s+led\s+by|led\s+by"
    r"|point\s+of\s+contact\s+(?:for\s+this\s+review\s+)?is"
    r"|prepared\s+by|handled\s+by|conducted\s+by|escalated\s+by"
    r"|interviewed\s+by|contact\s+(?:for\s+this\s+review\s+)?is"
    r")"
    # Allow optional comma or colon between keyword and name (e.g. "officer, Neha Joshi")
    r"[,:\s]+"
    r"([A-Z][a-z]{1,20}\s+[A-Z][a-z]{1,20})"
    r"(?=\s*[\(\,\.\n\r]"
    r"|\s+(?:from|via|using|at|regarding|contacted|email|phone|\(|IP|"
    r"on|noted|handled|confirmed|initiated|escalated|denied|can|after|who|"
    r"reported|reachable|and)\b)",
    re.IGNORECASE
)

NON_NAME_WORDS = {
    "activity","trends","success","teams","support","operations","summary","during",
    "archival","references","interaction","session","profile","feedback","engagement",
    "following","gathered","latest","product","update","service","request","portal",
    "engine","billing","internal","analytics","dashboard","tracking","feature",
    "adoption","customer","user","client","data","review","process","system",
    "compliance","infrastructure","quarterly","marketing","engineering","finance",
    "security","verified","identity","agent","case","account","description",
}

def extract_named_persons(text):
    names = set()
    for m in CUSTOMER_NAME_RE.finditer(text):
        candidate = m.group(1).strip()
        if not any(w in NON_NAME_WORDS for w in candidate.lower().split()):
            names.add(candidate)
    return names


def gazetteer_find_all_names(text):
    """Extract all names found by the gazetteer (for pre-seeding the vault)."""
    if not NER_AVAILABLE:
        return set()
    try:
        from ner_detect import gazetteer_find_names
        return {matched for (_, _, matched) in gazetteer_find_names(text)}
    except Exception:
        return set()


def preseed_vault(all_names):
    """
    Assign research-mode tokens to all known names before redaction starts.
    This guarantees token assignment order is deterministic (alphabetical)
    and prevents collisions between context-scanner names and gazetteer names.
    Names that appear in multiple files always get the same token.
    """
    for name in sorted(all_names):
        VAULT.get("name", name)


def auto_scan(folder_path, schema):
    print(col("info", "\n  ⟳ Auto-scanning files for PII patterns..."))

    if NER_AVAILABLE:
        loaded = load_spacy_model()
        if not loaded:
            print_setup()
    else:
        print(col("warn", "  ⚠  ner_detect.py not found — NER pass disabled."))

    all_names = set()
    unknown_fields = set()
    FIELD_HINT_RE = re.compile(
        r"(?:IFSC|SWIFT|VPA|UPI|IBAN|NPI|DEA|DL|License|Licence|"
        r"Reg(?:istration)?|MRN|PatientID|EmployeeID|StaffID)[:\s#]+([A-Z0-9\-]+)",
        re.IGNORECASE
    )

    for path in sorted(folder_path.iterdir()):
        if path.suffix.lower() not in READERS:
            continue
        text = READERS[path.suffix.lower()](path)
        if not text:
            continue
        all_names.update(extract_named_persons(text))
        all_names.update(gazetteer_find_all_names(text))
        for m in FIELD_HINT_RE.finditer(text):
            ft = m.group(0).split()[0].upper().rstrip(":#")
            if ft not in schema:
                unknown_fields.add(ft)

    if unknown_fields:
        print(col("warn", f"\n  ⚠  Unknown field types detected: {', '.join(sorted(unknown_fields))}"))
        for field in sorted(unknown_fields):
            if field in schema:
                continue
            answer = input(col("bold", f"\n  Is '{field}' PII? (y/n): ")).strip().lower()
            if answer == "y":
                modes_in = input(col("bold", "  Modes? (all / public,research,audit): ")).strip().lower()
                ml = ["public","research","audit"] if modes_in == "all" else [m.strip() for m in modes_in.split(",")]
                schema[field] = {"is_pii": True, "modes": ml}
            else:
                schema[field] = {"is_pii": False, "modes": []}
        save_schema(schema)
        print(col("keep", "  ✓ Schema saved to pii_schema.json"))

    # Pre-seed vault with all names in sorted order so tokens are
    # deterministic and collision-free across all detection passes
    preseed_vault(all_names)

    print(col("keep", f"  ✓ Scan complete. Found {len(all_names)} named person(s)."))
    return schema, all_names


# ─────────────────────────────────────────────────────────────
# REDACTION ENGINE
# ─────────────────────────────────────────────────────────────

def redact_text(text, mode, known_names, schema):
    stats = {}
    # Strip invisible unicode characters (braille blanks, zero-width spaces, etc.)
    # that appear in malformed data and break PAN/Aadhaar regex matching
    result = re.sub(r'[⠀-⣿​-‏﻿ ]', ' ', text)

    # Pass 1 — Regex structured PII
    for (pii_type, pattern, rp, rr, ra) in PII_RULES:
        if not {"public": rp, "research": rr, "audit": ra}[mode]:
            continue
        matches = re.findall(pattern, result, flags=re.IGNORECASE)
        if matches:
            stats[pii_type] = stats.get(pii_type, 0) + len(matches)
            result = re.sub(pattern,
                            lambda m, t=pii_type: replace(t, m.group(), mode),
                            result, flags=re.IGNORECASE)

    # Pass 2 — Context-scanned names (all modes — public/audit block, research pseudonymises)
    for name in known_names:
        pattern = re.escape(name)
        count = len(re.findall(pattern, result, flags=re.IGNORECASE))
        if count:
            stats["name"] = stats.get("name", 0) + count
            result = re.sub(pattern,
                            lambda m: replace("name", m.group(), mode),
                            result, flags=re.IGNORECASE)

    # Pass 3 — Schema-learned unknown fields
    for field, info in schema.items():
        if not info.get("is_pii") or mode not in info.get("modes", []):
            continue
        fp = rf"{re.escape(field)}[:\s#]+([A-Z0-9\-]+)"
        matches = re.findall(fp, result, re.IGNORECASE)
        if matches:
            stats[field] = stats.get(field, 0) + len(matches)
            result = re.sub(fp,
                            lambda m: m.group(0).replace(m.group(1),
                                                         replace("generic", m.group(1), mode)),
                            result, flags=re.IGNORECASE)

    # Pass 4 — NER (spaCy + Indian gazetteer + Stanza Hindi)
    if NER_AVAILABLE:
        result, ner_stats = ner_redact(result, mode)
        for k, v in ner_stats.items():
            stats[f"NER-{k}"] = stats.get(f"NER-{k}", 0) + v

    return result, stats


# ─────────────────────────────────────────────────────────────
# FILE WRITERS
# ─────────────────────────────────────────────────────────────

def write_txt(text, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

def write_pdf(text, out_path, original_path=None, mode=None, known_names=None):
    """In-place PDF redaction — removes text from content stream + draws black rectangles."""
    if original_path is None:
        write_txt(text, out_path.with_suffix(".txt"))
        return {}
    try:
        from pdf_inplace import redact_pdf_inplace
        return redact_pdf_inplace(original_path, out_path, mode, known_names or set())
    except Exception as e:
        print(col("warn", f"  [!] In-place PDF failed ({e}), saving as .txt"))
        write_txt(text, out_path.with_suffix(".txt"))
        return {}

def write_docx(text, out_path):
    from docx import Document
    doc = Document()
    c = doc.core_properties
    c.author = ""; c.last_modified_by = ""; c.title = ""; c.subject = ""; c.keywords = ""
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(out_path))

def write_csv(text, out_path):
    with open(out_path, "w", encoding="utf-8", newline="") as f:
        f.write(text)

def write_xlsx(original_path, redacted_text, out_path):
    try:
        import openpyxl
        wb_orig = openpyxl.load_workbook(str(original_path), data_only=True)
        wb_new  = openpyxl.Workbook()
        wb_new.remove(wb_new.active)
        lines = redacted_text.split("\n")
        li = 0
        for ws_orig in wb_orig.worksheets:
            ws_new = wb_new.create_sheet(title=ws_orig.title)
            for row in ws_orig.iter_rows():
                cells = lines[li].split("\t") if li < len(lines) else []
                li += 1
                for ci, cell in enumerate(row):
                    ws_new.cell(row=cell.row, column=cell.column,
                                value=cells[ci] if ci < len(cells) else "")
        wb_new.properties.creator = ""; wb_new.properties.lastModifiedBy = ""
        wb_new.properties.title = ""
        wb_new.save(str(out_path))
    except Exception as e:
        write_txt(redacted_text, out_path.with_suffix(".txt"))
        print(col("warn", f"  [!] XLSX write failed ({e}), saved as .txt"))

def write_sql(text, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

WRITERS = {
    ".txt":  lambda t, op, orig, **kw: write_txt(t, op),
    ".pdf":  lambda t, op, orig, **kw: write_pdf(t, op, original_path=orig,
                                                   mode=kw.get("mode"),
                                                   known_names=kw.get("known_names")),
    ".docx": lambda t, op, orig, **kw: write_docx(t, op),
    ".csv":  lambda t, op, orig, **kw: write_csv(t, op),
    ".xlsx": lambda t, op, orig, **kw: write_xlsx(orig, t, op),
    ".sql":  lambda t, op, orig, **kw: write_sql(t, op),
}


# ─────────────────────────────────────────────────────────────
# AUDIT LOG
# ─────────────────────────────────────────────────────────────

def write_log(log_rows, out_folder):
    log_path = out_folder / "redaction_log.csv"
    with open(log_path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=[
            "timestamp","original_file","output_file","mode",
            "pii_types_redacted","total_redactions"])
        w.writeheader()
        w.writerows(log_rows)
    print(col("keep", f"\n  ✓ Log saved → {log_path}"))


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Offline PII Redactor")
    parser.add_argument("folder", help="Folder containing files to redact")
    parser.add_argument("--mode", required=True, choices=["public","research","audit"])
    args = parser.parse_args()

    folder = Path(args.folder)
    if not folder.is_dir():
        print(col("redact", f"[!] Not a directory: {folder}")); sys.exit(1)

    mode = args.mode
    out_folder = folder / "redacted"
    out_folder.mkdir(exist_ok=True)

    banner(f"PII REDACTOR  |  Mode: {mode.upper()}")
    print(col("dim",  f"  Input  : {folder}"))
    print(col("dim",  f"  Output : {out_folder}"))
    print(col("info", f"  Strategy : " + {
        "public":   "Fixed uniform blocks ████████████ — zero length info",
        "research": "Consistent pseudonyms [NAME-R001] — re-linkable, not re-identifiable",
        "audit":    "Type labels [NAME][EMAIL] — identity stripped, operations intact",
    }[mode]))

    schema = load_schema()
    schema, known_names = auto_scan(folder, schema)
    if known_names:
        print(col("dim", f"  Named persons : {', '.join(sorted(known_names))}"))

    files = sorted([p for p in folder.iterdir()
                    if p.suffix.lower() in READERS and p.is_file()])
    if not files:
        print(col("warn", "\n  No supported files found.")); sys.exit(0)

    print(col("info", f"\n  Processing {len(files)} file(s)...\n"))
    log_rows = []

    for path in files:
        ext = path.suffix.lower()
        print(col("bold", f"  ▶ {path.name}"))

        text = READERS[ext](path)
        if not text.strip():
            print(col("dim", "    (empty or unreadable — skipped)\n")); continue

        redacted, stats = redact_text(text, mode, known_names, schema)
        total = sum(stats.values())

        if stats:
            for pt, count in sorted(stats.items()):
                print(col("redact", f"    ✂  {pt}: {count} instance(s)"))
        else:
            print(col("keep", "    ✓  No PII found"))

        out_name = f"{path.stem}_{mode}{ext}"
        out_path = out_folder / out_name
        pdf_stats = WRITERS[ext](redacted, out_path, path, mode=mode, known_names=known_names)

        if ext == ".pdf" and pdf_stats:
            stats = pdf_stats
            total = sum(stats.values())
            for pt, count in sorted(stats.items()):
                print(col("redact", f"    ✂  {pt}: {count} [PDF stream]"))

        print(col("keep", f"    → Saved: {out_name}\n"))
        log_rows.append({
            "timestamp":          datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "original_file":      path.name,
            "output_file":        out_name,
            "mode":               mode,
            "pii_types_redacted": ", ".join(sorted(stats.keys())) if stats else "none",
            "total_redactions":   total,
        })

    write_log(log_rows, out_folder)

    if mode == "research":
        summary = VAULT.summary()
        if summary:
            print(col("info", "\n  Pseudonym vault — unique values tokenised this run:"))
            for t, n in sorted(summary.items()):
                print(col("dim", f"    {t}: {n} unique value(s)"))
            print(col("warn", "  Tokens reset on next run (in-memory only, not saved to disk)."))

    total_all = sum(r["total_redactions"] for r in log_rows)
    banner(f"DONE  |  {len(files)} files  |  {total_all} total redactions")
    print(col("dim", f"  Redacted files in: {out_folder}\n"))


if __name__ == "__main__":
    main()
